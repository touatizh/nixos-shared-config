from docx import Document
import sys
import os


def md_to_docx(input_md: str, output_docx: str) -> None:
    # Create a new document
    doc: Document = Document()

    # Document title
    doc.add_heading("UH1/AB205 Tasks Standards (Tasks 1000-2169)", level=1)

    # Read markdown file
    with open(input_md, "r", encoding="utf-8") as file:
        content: str = file.read()

    # Split tasks by "---"
    tasks: list[str] = content.split("\n---\n")

    # Process individual tasks
    for task in tasks:
        lines: list[str] = task.strip().split("\n")
        if lines:
            # Task title
            doc.add_heading(lines[0].strip().split("### ")[1], level=2)
            # Task details
            for line in lines[1:]:
                if "**" in line.strip():
                    doc.add_heading(line.strip().split("**")[1], level=3)
                else:
                    doc.add_paragraph(line.strip())

    doc.save(output_docx)
    print(f"Successfully converted. Output file: {output_docx}")


if __name__ == "__main__":
    # Check for args
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input_markdown_file> <output_docx_file>")
        sys.exit(1)

    input_md: str = sys.argv[1]
    output_docx: str = sys.argv[2]

    # Check for input file
    if not os.path.isfile(input_md):
        print(f"Error: Input file '{input_md}' does not exist.")
        sys.exit(1)

    # Generate the DOCX file
    md_to_docx(input_md, output_docx)
